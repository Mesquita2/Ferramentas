from datetime import date
import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.shared import Inches
import streamlit as st
import pandas as pd
import io

imagem_rodape = "Endereço.jpeg"
imagem_cabecalho = 'Logo.jpg'
ARQUIVOBASE = "alunosxdisciplinas"
ARQUIVOREC = "rec"


def carregar():
        
    def gerar_relatorio_nao_encontrados(df):
        if df.empty:
            return None

        doc = Document()
        
        # Título
        doc.add_heading('Alunos Removidos - Fora do Período em Curso', level=1)

        # Subtítulo com quantidade
        doc.add_paragraph(f"Total de alunos removidos: {len(df)}\n")

        # Tabela
        tabela = doc.add_table(rows=1, cols=3)
        tabela.style = 'Table Grid'

        # Cabeçalho
        hdr_cells = tabela.rows[0].cells
        hdr_cells[0].text = 'RA'
        hdr_cells[1].text = 'Nome'
        hdr_cells[2].text = 'Disciplina'

        # Dados
        for _, row in df.iterrows():
            row_cells = tabela.add_row().cells
            row_cells[0].text = str(row['RA'])
            row_cells[1].text = str(row['NOME'])
            row_cells[2].text = str(row['DISCIPLINA'])

        # Salvar em memória
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        return buffer


    def limpar_rec(df_rec):
        if df_rec is not None:
            df_base = st.session_state["dados"].get(ARQUIVOBASE).copy()

            df_rec['DISCIPLINA'] = (
                df_rec['DISCIPLINA']
                .str.replace(r'\s*\([^()]*\)\s*$', '', regex=True)
                .str.replace(r'[\u200b\u200e\u202c\u00a0]', '', regex=True)
                .str.strip()
            )

            df_rec["RA"] = df_rec["RA"].astype(str).str.zfill(7)
            df_base["RA"] = df_base["RA"].astype(str).str.zfill(7)

            df_base.rename(columns={'NOMEDISCIPLINA': 'DISCIPLINA',
                                    'NOMEALUNO': 'ALUNO'}, inplace=True)

            # Filtrar somente alunos com "Período em Curso"
            df_base = df_base[df_base['NOMESTATUS'] == 'Período em Curso']

            # Merge para trazer TURMADISC e validar se está em curso
            df_merged = pd.merge(
                df_rec,
                df_base[['RA', 'DISCIPLINA', 'TURMADISC', 'ALUNO']],
                on=['RA', 'DISCIPLINA'],
                how='inner'  # só mantém quem está nas duas bases
            )

            df_merged = df_merged.drop_duplicates(subset=['RA', 'DISCIPLINA', 'TURMADISC'])
            df_merged = df_merged[df_merged['RA'].notna()]
            
            nao_encontrados = df_rec.merge(df_merged[['RA', 'DISCIPLINA']], 
                                        on=['RA', 'DISCIPLINA'], 
                                        how='left', indicator=True)
            nao_encontrados = nao_encontrados[nao_encontrados['_merge'] == 'left_only']

            if not nao_encontrados.empty:
                st.warning(f"{len(nao_encontrados)} alunos não estão em 'Período em Curso' e foram removidos:")
                st.dataframe(nao_encontrados[['RA', 'NOME', 'DISCIPLINA']])
                
            buffer = gerar_relatorio_nao_encontrados(nao_encontrados)

            if buffer:
                st.download_button(
                    label="📄 Baixar Relatório Word",
                    data=buffer,
                    file_name="alunos_removidos.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

            if df_merged.empty:
                st.warning("Nenhum aluno inscrito está com status 'Período em Curso'.")
            else:
                st.success("Alunos em 'Período em Curso' carregados com sucesso!")
            return df_merged
        else:
            st.warning("Arquivo REC não encontrado.")
            return pd.DataFrame()
            
    def adicionar_imagem_no_cabecalho(doc, imagem_cabecalho):
        # Acessando o cabeçalho da primeira seção do documento
        section = doc.sections[0]
        header = section.header

        # Criando um parágrafo no cabeçalho e adicionando uma imagem
        paragraph = header.paragraphs[0]  # Usando o primeiro parágrafo do cabeçalho
        section.header_distance = Inches(0.2)   
        run = paragraph.add_run()

        # Adicionando a imagem ao cabeçalho
        run.add_picture(imagem_cabecalho, width=Inches(7.5), height=Inches(1))  # Ajuste o tamanho conforme necessário

    def adicionar_imagem_no_rodape(doc, imagem_rodape):
        # Acessando o rodapé da primeira seção do documento
        section = doc.sections[0]
        footer = section.footer
        

        # Criando um parágrafo no rodapé e adicionando uma imagem
        paragraph = footer.paragraphs[0]  # Usando o primeiro parágrafo do rodapé
        section.footer_distance = Inches(0.2)
        run = paragraph.add_run()

        # Adicionando a imagem ao rodapé
        run.add_picture(imagem_rodape, width=Inches(7.5), height=Inches(1))  # Ajuste o tamanho conforme necessário


    def gerar_relatorio(df, disciplinas, turmas):
        dataatual = date.today().strftime('%d/%m/%Y')
        df = df[df["DISCIPLINA"].isin(disciplinas) & df["TURMADISC"].isin(turmas)]
        df = df.sort_values(by=["DISCIPLINA", "TURMADISC", "ALUNO"])

        doc = Document()
        adicionar_imagem_no_cabecalho(doc, imagem_cabecalho)
        adicionar_imagem_no_rodape(doc, imagem_rodape)

        section = doc.sections[0]
        section.left_margin = Inches(0.5) 
        section.right_margin = Inches(0.5)
        section.top_margin = Inches(0.5)  
        section.bottom_margin = Inches(0.5) 

        for (disciplina, turma), df_grupo in df.groupby(['DISCIPLINA', 'TURMADISC']):
            # Título
            p = doc.add_paragraph()
            p.add_run("\n\n")
            run = p.add_run(f"Disciplina: {disciplina}")
            run.font.name = 'Arial'
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0, 0, 0)

            p = doc.add_paragraph()
            run = p.add_run(f"Turma: {turma}")
            run.font.name = 'Arial'
            run.font.size = Pt(12)

            p = doc.add_paragraph()
            run = p.add_run(f"Data: {dataatual}")
            run.font.name = 'Arial'
            run.font.size = Pt(12)

            # Tabela
            df_grupo = df_grupo[['ALUNO']].copy()
            df_grupo['ASSINATURA'] = ''
            table = doc.add_table(rows=1, cols=len(df_grupo.columns))
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            for i, col_name in enumerate(df_grupo.columns):
                hdr_cells[i].text = col_name

            for _, row in df_grupo.iterrows():
                row_cells = table.add_row().cells
                for i, item in enumerate(row):
                    row_cells[i].text = str(item)

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    def dash(df):
        if not df:
            st.write("Data frame Vazio")
            return pd.DataFrame() 
        if not os.path.exists(df):
            st.write(f"Erro: Arquivo '{df}' não encontrado.")
            return pd.DataFrame()  
        return pd.read_excel(df) 

    # Interface do Streamlit
    st.title("Limpeza e tratamento de notas de REC")
            
    st.subheader("Dados dos Cadastrados na REC")
    df_cadastro = st.session_state["dados"].get(ARQUIVOREC).copy()
    df = df_cadastro.copy()
    if df_cadastro is not None: 
        st.dataframe(df_cadastro[['DISCIPLINA', 'NOME']])


    def gerar_excel(df_rec, disciplinas, turmas):
        df_filtrado = df_rec[df_rec["DISCIPLINA"].isin(disciplinas) & df_rec["TURMADISC"].isin(turmas)].copy()
        df_filtrado['RA'] = df_filtrado['RA'].astype(str).str.zfill(7)
        df_filtrado['NOTAS'] = 0
        colunas = ['TURMADISC', 'DISCIPLINA', 'RA', 'ALUNO', 'NOTAS']
        df_filtrado = df_filtrado[colunas].sort_values(by=["DISCIPLINA", "TURMADISC", "ALUNO"])
        
        output = io.BytesIO()
        with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
            df_filtrado.to_excel(writer, index=False, sheet_name="Notas")
        output.seek(0)
        return output



    st.title("Gerador de Planilha de Notas para REC")

    df_rec = limpar_rec(df)
    if df_rec.empty:
        st.stop()
            
    disciplinas = df_rec["DISCIPLINA"].unique().tolist()
    disciplinas_selecionadas = st.multiselect("1. Escolha as disciplinas", disciplinas)

    if disciplinas_selecionadas:
        turmas_disponiveis = df_rec[df_rec["DISCIPLINA"].isin(disciplinas_selecionadas)]["TURMADISC"].unique().tolist()
        turmas_selecionadas = st.multiselect("2. Escolha as turmas", turmas_disponiveis)

        if turmas_selecionadas:
            prova = st.selectbox("3. Escolha se é REC_P1 ou REC_P2 ou REC_FINAL", ["REC_P1", "REC_P2", "REC_FINAL"])

            # Filtra os dados para visualização
            df_filtrado = df_rec[
                (df_rec["DISCIPLINA"].isin(disciplinas_selecionadas)) &
                (df_rec["TURMADISC"].isin(turmas_selecionadas))
            ]

            st.write(f"**Alunos da(s) Disciplina(s): {disciplinas_selecionadas} | Turma(s): {turmas_selecionadas}**")
            total = df_filtrado['ALUNO'].count()
            st.write(f"**Quantidade de REC solicitadas: {total}**")
            qtd_aluno = df_filtrado['ALUNO'].drop_duplicates().count()
            st.write(f"**Quantidade de alunos: {qtd_aluno}**")
            df_filtrado = df_filtrado.sort_values(by="ALUNO", ascending=True)
            st.dataframe(df_filtrado[["ALUNO", "DISCIPLINA", "TURMADISC"]])

            # Botão para gerar planilha Excel
            excel_file = gerar_excel(df_rec, disciplinas_selecionadas, turmas_selecionadas)
            st.download_button(
                label="⬇ Gerar e Baixar Planilha Excel (Multi)",
                data=excel_file,
                file_name=f"Planilha_REC_{prova}.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            )

            st.title("Criar Relatório de Assinatura")
            relatorio = gerar_relatorio(df_rec, disciplinas_selecionadas, turmas_selecionadas)
            st.download_button(
                label="📄 Gerar e Baixar Relatório de Assinaturas",
                data=relatorio,
                file_name=f"Relatorio_Assinatura_REC_{prova}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

