from datetime import date
import io
import pandas as pd
import streamlit as st
from docx import Document
from docx.shared import Inches
import os
from docx.enum.text import WD_ALIGN_PARAGRAPH

def carregar():
    imagem_rodape = "./Endereço.jpeg"
    imagem_cabecalho = "./Logo.jpg"
    ARQUIVOBASE = "alunosxdisciplinas_geral"
    
    import unicodedata

    def normalizar_texto(s):
        s = str(s).strip().lower()
        s = unicodedata.normalize("NFKD", s)
        s = "".join(c for c in s if not unicodedata.combining(c))
        return s

    def gerar_relatorio_assinatura(df_alunos, curso, periodo, turma, data_hoje, imagem_cabecalho, imagem_rodape):
        data_hoje = data_hoje.strftime("%d/%m/%Y")
        df_sorted = df_alunos.sort_values(by=["ALUNO"]).reset_index(drop=True)

        doc = Document()
        section = doc.sections[0]
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)

        # Cabeçalho
        header = section.header
        header_paragraph = header.paragraphs[0]
        header_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = header_paragraph.add_run()
        if os.path.exists(imagem_cabecalho):
            run.add_picture(imagem_cabecalho, width=Inches(8))

        # Rodapé
        footer = section.footer
        footer_paragraph = footer.paragraphs[0]
        footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_footer = footer_paragraph.add_run()
        if os.path.exists(imagem_rodape):
            run_footer.add_picture(imagem_rodape, width=Inches(8))

        # Info turma
        if curso: doc.add_paragraph(f"Curso: {curso}", style='Heading 2')
        if turma: doc.add_paragraph(f"Turma: {turma}")
        if periodo: doc.add_paragraph(f"Período: {periodo}")
        doc.add_paragraph(f"Data: {data_hoje}")
        doc.add_paragraph(" ")

        # Tabela
        tabela = doc.add_table(rows=1, cols=4)
        tabela.style = "Table Grid"
        hdr_cells = tabela.rows[0].cells
        hdr_cells[0].text = 'Aluno'
        hdr_cells[3].text = 'Assinatura'

        for _, row in df_sorted.iterrows():
            linha = tabela.add_row().cells
            linha[0].text = str(row.get("ALUNO", ""))[:200]
            linha[3].text = " "

        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        return output
    
    def sigla_curso(nome):
        palavras = nome.split()
        return "".join(p[0].upper() for p in palavras[:3])  # até 3 letras (ajuste se quiser mais)
    
    def gerar_excel_com_filtros(df_rec):
        df_filtrado = df_rec.copy()
        output = io.BytesIO()
        with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
            df_filtrado.to_excel(writer, index=False, sheet_name="Notas")
        output.seek(0)
        return output

    # Função para converter número em romano
    def para_romano(num):
        romanos = {1:'I', 2:'II', 3:'III', 4:'IV', 5:'V', 6:'VI', 7:'VII', 8:'VIII', 9:'IX', 10:'X'}
        return romanos.get(num, str(num))  # fallback p/ número normal se > 10
    
    def ajustes_dataframe(df):
        """
        Normaliza o DataFrame de upload (ZipGrade/Excel).
        - Garante coluna 'RA' como string com 7 dígitos (zero-padded).
        - Garante coluna 'NOMEALUNO' com nome completo do aluno.
        - Remove linhas com RA igual a '0' ou nome vazio.
        - Retorna uma cópia do DataFrame com as colunas originais preservadas
        (apenas adiciona/renomeia 'RA' e 'NOMEALUNO' quando necessário).
        """
        df = df.copy()

        # --- RA ---
        # Prioriza coluna já existente 'RA', senão tenta 'Student ID' ou 'StudentID'
        if 'RA' in df.columns:
            df['RA'] = df['RA'].astype(str).str.strip()
        elif 'Student ID' in df.columns:
            df['RA'] = df['Student ID'].astype(str).str.strip()
        elif 'StudentID' in df.columns:
            df['RA'] = df['StudentID'].astype(str).str.strip()
        else:
            # cria coluna vazia para evitar erros posteriores
            df['RA'] = ''

        # Remove pontos/virgulas e zfill
        df['RA'] = df['RA'].replace({'nan': '', 'None': ''}).fillna('').astype(str)
        # se houver valores com casas decimais como '12345.0', remove parte decimal
        df['RA'] = df['RA'].apply(lambda x: x.split('.')[0] if isinstance(x, str) and '.' in x else x)
        df['RA'] = df['RA'].str.replace(r'\D', '', regex=True)  # remove não dígitos
        df['RA'] = df['RA'].apply(lambda x: x.zfill(7) if x != '' else '')

        # --- NOMEALUNO ---
        if 'NOMEALUNO' in df.columns:
            df['NOMEALUNO'] = df['NOMEALUNO'].astype(str).str.strip()
        else:
            first_name_cols = ['Student First Name', 'First Name', 'Nome', 'Nome completo', 'Name']
            last_name_cols = ['Student Last Name', 'Last Name', 'Sobrenome', 'Surname']

            first = None
            last = None

            for c in first_name_cols:
                if c in df.columns:
                    first = c
                    break
            for c in last_name_cols:
                if c in df.columns:
                    last = c
                    break

            if first and last:
                df['NOMEALUNO'] = (df[first].fillna('').astype(str) + ' ' + df[last].fillna('').astype(str)).str.strip()
            elif 'Nome completo' in df.columns:
                df['NOMEALUNO'] = df['Nome completo'].fillna('').astype(str).str.strip()
            elif 'Nome' in df.columns:
                df['NOMEALUNO'] = df['Nome'].fillna('').astype(str).str.strip()
            elif 'Student Name' in df.columns:
                df['NOMEALUNO'] = df['Student Name'].fillna('').astype(str).str.strip()
            else:
                # última tentativa: concatena quaisquer colunas que pareçam nome
                possible_name_cols = [c for c in df.columns if 'name' in c.lower() or 'nome' in c.lower()]
                if possible_name_cols:
                    df['NOMEALUNO'] = df[possible_name_cols].astype(str).agg(' '.join, axis=1).str.replace(r'\s+', ' ', regex=True).str.strip()
                else:
                    df['NOMEALUNO'] = ''

        # --- Filtrar linhas inválidas ---
        # RA inválido (vazio ou '0000000' ou '0' padronizado para 7) e nome vazio
        df['RA_clean_flag'] = df['RA'].apply(lambda x: (not x) or x == '0000000')
        df['NOMEALUNO'] = df['NOMEALUNO'].fillna('').astype(str)
        valid_mask = (~df['RA_clean_flag']) & (df['NOMEALUNO'] != '')

        df = df.loc[valid_mask].copy()
        df.drop(columns=['RA_clean_flag'], inplace=True)

        # --- Garante tipos e zfill final (por segurança) ---
        df['RA'] = df['RA'].astype(str).str.zfill(7)

        # Mantém todas as outras colunas; somente assegura RA e NOMEALUNO presentes
        return df
    
    st.title("Sistema de Nivelamento de Inglês")
    
    tab1, tab2, tab3= st.tabs(["Relatório de Alunos", "Notas do Nivelamento", "Formatar para Questionário"])
    
    # ---------------------------
    # TAB 1: relatório de alunos (corrigido)
    # ---------------------------
    with tab1:
        st.subheader("Upload do arquivo de Nivelamento de Inglês")
        # uploader específico para o tab1 com key único
        uploaded_file_rel = st.file_uploader(
            "Selecione o arquivo (Excel ou CSV) — Relatório de Alunos",
            type=["xlsx", "csv"],
            key="relatorio"
        )

        if uploaded_file_rel is None:
            st.info("Envie um arquivo para começar (aba: Relatório de Alunos).")
        else:
            # le o arquivo com tratamento
            try:
                if str(uploaded_file_rel.name).lower().endswith(".xlsx"):
                    df_env = pd.read_excel(uploaded_file_rel)
                else:
                    # tenta ; primeiro, se falhar tenta ,
                    try:
                        df_env = pd.read_csv(uploaded_file_rel, sep=";")
                    except Exception:
                        df_env = pd.read_csv(uploaded_file_rel, sep=",")
            except Exception as e:
                st.error(f"Erro ao ler o arquivo enviado: {e}")
                df_env = None

            # se a leitura falhou, não tenta processar
            if df_env is None:
                st.warning("Não foi possível processar o arquivo. Verifique o formato e tente novamente.")
            else:
                # --- normaliza colunas com segurança ---
                df_env.rename(columns={
                    'Curso': 'CURSO', 'Nome completo': 'ALUNO', 'Nome': 'ALUNO',
                    'Período atual': 'PERIODO', 'Período': 'PERIODO',
                    'E-mail institucional': 'EMAIL', 'E-mail': 'EMAIL',
                    'Email': 'EMAIL', 'email': 'EMAIL', 'RA': 'RA', 'Turma': 'TURMA'
                }, inplace=True, errors="ignore")

                # garante colunas mínimas
                for col in ['EMAIL','PERIODO']:
                    if col not in df_env.columns:
                        df_env[col] = ''

                # filtra só registros com email (como você fazia)
                df_env = df_env[df_env['EMAIL'].notna() & (df_env['EMAIL'] != '')].copy()

                # --- prepara df_base local (ja validado antes dos tabs) ---
                # assume que df_base foi carregado e padronizado antes (veja seu fluxo)
                df_base_local = st.session_state["dados"].get(ARQUIVOBASE).copy()
                df_base_local.rename(columns={
                    'Curso': 'CURSO','Aluno': 'ALUNO','Nome completo': 'ALUNO',
                    'Período atual': 'PERIODO','Período': 'PERIODO',
                    'E-mail institucional': 'EMAIL','E-mail': 'EMAIL','Email': 'EMAIL','email': 'EMAIL',
                    'RA': 'RA','Turma': 'TURMA'
                }, inplace=True, errors="ignore")

                if 'RA' in df_base_local.columns:
                    df_base_local["RA"] = df_base_local["RA"].astype(str).str.zfill(7)
                for col in ['ALUNO', 'RA', 'CURSO', 'EMAIL', 'TURMA', 'PERIODO']:
                    if col not in df_base_local.columns:
                        df_base_local[col] = ''
                df_base_local = df_base_local[['ALUNO','RA','CURSO','EMAIL','TURMA','PERIODO']].copy()

                # atualiza PERIODO via envio (se existir)
                try:
                    df_base_local = df_base_local.merge(
                        df_env[['EMAIL','PERIODO']].drop_duplicates(subset=['EMAIL', 'PERIODO']),
                        on='EMAIL', how='left', suffixes=('','_env')
                    )
                    if 'PERIODO_env' in df_base_local.columns:
                        df_base_local['PERIODO'] = df_base_local['PERIODO_env'].fillna(df_base_local['PERIODO'])
                        df_base_local.drop(columns=['PERIODO_env'], inplace=True)
                except Exception as e:
                    st.warning(f"Falha ao atualizar período a partir do arquivo enviado: {e}")

                # cruzamento por EMAIL
                try:
                    df_cruzado = df_base_local[df_base_local['EMAIL'].isin(df_env['EMAIL'])].copy()
                    df_nao_encontrados = df_env[~df_env['EMAIL'].isin(df_base_local['EMAIL'])].copy()
                except Exception as e:
                    st.error(f"Erro ao cruzar bases: {e}")
                    df_cruzado = pd.DataFrame()
                    df_nao_encontrados = pd.DataFrame()

                # exibe encontrados
                if not df_cruzado.empty:
                    st.subheader("Alunos encontrados no ARQUIVOBASE")
                    st.dataframe(df_cruzado[['ALUNO','RA','EMAIL','CURSO','TURMA','PERIODO']])

                    periodo = df_cruzado['PERIODO'].dropna().unique().tolist()
                    periodo_sel = None
                    if periodo:
                        periodo_sel = st.selectbox("Selecione a Turma/Período", ["Todos os Periodos"] + sorted(periodo), key="periodo_rel")
                        if periodo_sel == "Todos os Periodos":
                            periodo_sel = None

                    df_para_relatorio = df_cruzado.copy()
                    if periodo_sel:
                        df_para_relatorio = df_para_relatorio[df_para_relatorio['PERIODO'] == periodo_sel]

                    curso_head = df_para_relatorio['CURSO'].dropna().unique().tolist()
                    curso_head = curso_head[0] if curso_head else ''
                    periodo_head = ", ".join(df_para_relatorio['PERIODO'].dropna().unique().tolist())

                    # botão gerar docx
                    if st.button("Gerar relatório .docx desta Turma", key="btn_docx_rel"):
                        relatorio_docx = gerar_relatorio_assinatura(
                            df_para_relatorio, curso_head, periodo_head,
                            turma=(periodo_sel if periodo_sel else ''),
                            data_hoje=date.today(),
                            imagem_cabecalho=imagem_cabecalho,
                            imagem_rodape=imagem_rodape
                        )
                        st.download_button(
                            label="Download Relatório (.docx)",
                            data=relatorio_docx,
                            file_name=f"Relatorio_Assinaturas_{curso_head}_{periodo_sel or 'todas'}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                else:
                    st.info("Nenhum aluno encontrado na base para os emails do arquivo enviado.")

                # exibe não encontrados
                if not df_nao_encontrados.empty:
                    st.subheader("Alunos NÃO encontrados")
                    st.dataframe(df_nao_encontrados[['ALUNO','EMAIL','CURSO','PERIODO']])
                    if st.button("Gerar relatório .docx dos não encontrados", key="btn_docx_nao"):
                        relatorio_nao_encontrados = gerar_relatorio_assinatura(
                            df_nao_encontrados, curso='-', periodo='-', turma='-',
                            data_hoje=date.today(),
                            imagem_cabecalho=imagem_cabecalho,
                            imagem_rodape=imagem_rodape
                        )
                        st.download_button(
                            label="Download Relatório Não Encontrados (.docx)",
                            data=relatorio_nao_encontrados,
                            file_name="Relatorio_Assinaturas_Nao_Encontrados.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )

                # excel completo (duas abas) — só se houver dados
                if (not df_cruzado.empty) or (not df_nao_encontrados.empty):
                    import io  
                    output_excel_completo = io.BytesIO()
                    try:
                        with pd.ExcelWriter(output_excel_completo, engine='xlsxwriter') as writer:
                            if not df_cruzado.empty:
                                df_para_relatorio.to_excel(writer, sheet_name='Encontrados', index=False)
                            if not df_nao_encontrados.empty:
                                df_nao_encontrados.to_excel(writer, sheet_name='Nao_Encontrados', index=False)
                        output_excel_completo.seek(0)
                        st.download_button(
                            label="Download Excel Completo (Encontrados + Não Encontrados)",
                            data=output_excel_completo,
                            file_name=f"Relatorio_Completo_{curso_head}_{periodo_sel or 'todas'}.xlsx",
                            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                        )
                    except Exception as e:
                        st.warning(f"Não foi possível gerar o Excel completo: {e}")

    with tab2:
        import numpy as np  # import local para garantir disponibilidade
        st.subheader("Notas dos Alunos — cálculo direto a partir do arquivo")

        # tenta carregar a base de alunos (opcional)
        ARQUIVOBASE = "alunosxdisciplinas_geral"
        df_base = None
        if "dados" in st.session_state and isinstance(st.session_state["dados"], dict):
            df_base = st.session_state["dados"].get(ARQUIVOBASE)
        if df_base is None:
            df_base = st.session_state.get(ARQUIVOBASE)

        if df_base is None:
            st.warning("Base 'alunosxdisciplinas_geral' não encontrada no session_state — o cálculo seguirá, mas CURSO/TURMA não serão preenchidos.")
        else:
            # padroniza colunas importantes na base (se existir)
            df_base = df_base.copy()
            df_base.rename(columns={
                'Curso': 'CURSO',
                'Aluno': 'ALUNO',
                'Turma': 'TURMADISC'
            }, inplace=True, errors="ignore")
            for col in ['RA','ALUNO','CURSO','TURMADISC','IDTURMADISC','CODCOLIGADA']:
                if col not in df_base.columns:
                    df_base[col] = ''

            # criar mapeamento RA -> primeiro registro (evita duplicidade)
            df_base['RA'] = df_base['RA'].astype(str).str.replace(r'\D','', regex=True).str.zfill(7)
            mapa_alunos = df_base.drop_duplicates(subset=['RA']).set_index('RA')[['CODCOLIGADA','CURSO','TURMADISC','IDTURMADISC','ALUNO']]

        # uploader específico desta aba (usa key para não conflitar)
        uploaded_file = st.file_uploader("Envie o arquivo de notas (Excel/CSV) — apenas envie o export do ZipGrade/Sheet", type=["xlsx", "csv"], key="notas_tab")
        if uploaded_file is None:
            st.info("Envie um arquivo para processar (aba Notas).")
        else:
            # leitura com fallback de separadores
            try:
                if str(uploaded_file.name).lower().endswith(".xlsx"):
                    df_upload_raw = pd.read_excel(uploaded_file)
                else:
                    try:
                        df_upload_raw = pd.read_csv(uploaded_file, sep=";")
                    except Exception:
                        df_upload_raw = pd.read_csv(uploaded_file, sep=",")
            except Exception as e:
                st.error(f"Erro ao ler o arquivo enviado: {e}")
                df_upload_raw = None

            if df_upload_raw is None:
                st.warning("Não foi possível ler o arquivo enviado.")
            else:
                st.subheader("Preview (dados enviados)")
                st.dataframe(df_upload_raw.head(200))

                # normaliza/ajusta upload (usa sua função)
                df_upload = ajustes_dataframe(df_upload_raw)

                # Garante RA e NOMEALUNO
                if 'RA' not in df_upload.columns or df_upload['RA'].isna().all():
                    if 'Student ID' in df_upload_raw.columns:
                        df_upload['RA'] = df_upload_raw['Student ID'].astype(str).str.replace(r'\D','', regex=True).str.zfill(7)
                    else:
                        df_upload['RA'] = df_upload['RA'].astype(str).fillna('').str.replace(r'\D','', regex=True).str.zfill(7)
                df_upload['RA'] = df_upload['RA'].astype(str).str.zfill(7)
                if 'NOMEALUNO' not in df_upload.columns or df_upload['NOMEALUNO'].isna().all():
                    df_upload['NOMEALUNO'] = (
                        df_upload.get('Student First Name', '').fillna('') + ' ' +
                        df_upload.get('Student Last Name', '').fillna('')
                    ).str.strip()

                # Detecta colunas de resposta (#...) para contar NaNs (opcional)
                colunas_respostas = [c for c in df_upload_raw.columns if str(c).startswith("#")]
                if colunas_respostas:
                    df_upload['Nao_Respondidas'] = df_upload_raw[colunas_respostas].isna().sum(axis=1)
                    df_nulos = df_upload[df_upload['Nao_Respondidas'] > 0][['RA','NOMEALUNO','Nao_Respondidas']].copy()
                    if not df_nulos.empty:
                        st.subheader("Alunos com questões não respondidas")
                        st.dataframe(df_nulos.sort_values("Nao_Respondidas", ascending=False))
                        alunos_com_nulos = df_nulos['NOMEALUNO'].tolist()
                        selecionados = st.multiselect(
                            "Selecionar alunos para desconsiderar questões em branco (reduz Possible Points):",
                            options=alunos_com_nulos,
                            key="sel_nulos_notas"
                        )
                        alunos_ajustar = df_nulos[df_nulos['NOMEALUNO'].isin(selecionados)].set_index('RA')['Nao_Respondidas'].to_dict()
                    else:
                        alunos_ajustar = {}
                else:
                    alunos_ajustar = {}

                # Inputs de cálculo (simples)
                provas = [f'Nivelamento {romano}' for romano in ['I','II','III','IV','V','VI','VII','VIII']]
                prova = st.selectbox(
                    'Selecione o tipo de prova',
                    provas,
                    key="sel_tipo_prova_notas"
                )

                questoes_anuladas_input = st.text_input("Informe questões anuladas (separadas por vírgula):", value="", key="anuladas_notas")
                questoes_anuladas = [int(q.strip()) for q in questoes_anuladas_input.split(",") if q.strip().isdigit()]

                if st.button("Calcular Notas (apenas com o arquivo enviado)", key="btn_calcula_notas"):
                    # Verifica colunas necessárias
                    if 'Possible Points' not in df_upload.columns or 'Earned Points' not in df_upload.columns:
                        st.error("Arquivo inválido: verifique se existem as colunas 'Possible Points' e 'Earned Points' no arquivo exportado.")
                    else:
                        # Ajusta Possible Points para alunos selecionados (NaNs)
                        if alunos_ajustar:
                            df_upload['Possible Points Ajustado'] = df_upload.apply(
                                lambda row: row['Possible Points'] - alunos_ajustar.get(row['RA'], 0),
                                axis=1
                            )
                        else:
                            df_upload['Possible Points Ajustado'] = df_upload['Possible Points']

                        # Evita divisão por zero
                        df_upload['Possible Points Ajustado'] = df_upload['Possible Points Ajustado'].replace(0, np.nan)
                        df_upload['Earned Points Original'] = df_upload['Earned Points'].fillna(0)

                        # Bônus por questões anuladas (se existirem colunas #X Points Earned)
                        ids = df_upload['RA'].astype(str).str.zfill(7)
                        bonus_total = pd.Series(0, index=ids.unique())
                        for q in questoes_anuladas:
                            coluna = f"#{q} Points Earned"
                            if coluna in df_upload.columns:
                                ganhos = (df_upload[coluna].fillna(0) == 0).astype(int)
                                bonus = pd.Series(ganhos.values, index=ids).groupby(level=0).sum()
                                bonus_total = bonus_total.add(bonus, fill_value=0)

                        df_upload['Bonus Anuladas'] = ids.map(bonus_total).fillna(0)
                        df_upload['Earned Points Final'] = df_upload['Earned Points Original'] + df_upload['Bonus Anuladas']

                        # Regra de cálculo original
                        df_upload['NOTAS'] = np.minimum((df_upload['Earned Points Final']) / df_upload['Possible Points Ajustado'], 1).fillna(0) * 10                    
                        df_upload['STATUS_32'] = np.where(df_upload['Earned Points Final'] >= 32, 'Aprovado', 'Reprovado')
                        # Agrega por RA (em caso de múltiplas linhas por aluno no export)
                        # também traz primeiro NOMEALUNO caso exista
                        df_agregado = df_upload.groupby('RA', as_index=False).agg({
                            'NOTAS': 'mean',
                            'NOMEALUNO': 'first',
                            'STATUS_32': 'first',
                            'Earned Points Final': 'sum',
                        })

                        # Se tivermos a base de alunos, faz merge para trazer CURSO/TURMA/ID/etc
                        if df_base is not None:
                            # mapa_alunos foi criado antes (primeiro registro por RA)
                            df_agregado['RA'] = df_agregado['RA'].astype(str).str.zfill(7)
                            mapa_df = mapa_alunos.reset_index()
                            df_final = pd.merge(mapa_df, df_agregado, on='RA', how='right')
                            # Se mapa tiver colunas nulas, preenche com vazio
                            for c in ['CODCOLIGADA','CURSO','TURMADISC','IDTURMADISC','ALUNO']:
                                if c not in df_final.columns:
                                    df_final[c] = ''
                            # preferir NOMEALUNO do upload quando existir
                            df_final['ALUNO'] = df_final.apply(lambda r: r['NOMEALUNO'] if pd.notnull(r['NOMEALUNO']) and r['NOMEALUNO'] != '' else (r.get('ALUNO') or ''), axis=1)
                        else:
                            # sem base, monta colunas mínimas
                            df_final = df_agregado.copy()
                            df_final['ALUNO'] = df_final['NOMEALUNO']

                        # adiciona metadados da avaliação
                        df_final['PROVA'] = prova
                        df_final['STATUS_8'] = np.where(df_final['NOTAS'] >= 8, 'Aprovado', 'Reprovado')
                        

                        # organiza colunas no formato desejado
                        colunas = ['CURSO','TURMADISC','RA','ALUNO',
                                   'PROVA','NOTAS', 'STATUS_32', 'STATUS_8', 'Earned Points Final',]
                        # garante que existam todas as colunas
                        for c in colunas:
                            if c not in df_final.columns:
                                df_final[c] = ''

                        df_final = df_final[colunas]

                        # formata NOTAS com vírgula e 2 casas
                        df_final['NOTAS'] = pd.to_numeric(df_final['NOTAS'], errors='coerce').round(2)
                        df_final['NOTAS'] = df_final['NOTAS'].apply(lambda x: f"{x:.2f}".replace('.', ',') if pd.notnull(x) else '')

                        st.subheader("Notas calculadas (a partir do arquivo enviado)")
                        st.dataframe(df_final)

                        # Download em TXT (seu padrão)
                        import io 
                        output = io.BytesIO()
                        df_final.to_csv(output, index=False, sep=';', encoding='utf-8', header=False)
                        output.seek(0)
                        nome_arquivo = f"notas_enviadas_{prova}.txt"
                        st.download_button(
                            label="⬇ Baixar Notas Tratadas (TXT)",
                            data=output,
                            file_name=nome_arquivo,
                            mime="text/plain"
                        )
                        relatorio_excel = gerar_excel_com_filtros(df_final)
                        st.download_button(
                            label="Gerar Relatório de Notas",
                            data=relatorio_excel,
                            file_name=f"Relatorio_Notas_{prova}.xlsx",
                            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                        )

    with tab3:
        st.header("Formatar para Questionário (ZipGrade/Google Forms)")

        # tentar localizar df_base no session_state conforme seu padrão
        df_base = None
        if "dados" in st.session_state and isinstance(st.session_state["dados"], dict):
            df_base = st.session_state["dados"].get("alunosxdisciplinas")
            df_base = df_base.drop_duplicates(subset=['RA']) if df_base is not None else None
        if df_base is None:
            st.warning("df_base não encontrado em session_state. Carregue o banco 'alunosxdisciplinas' antes.")
            st.stop()

        st.write("Colunas do banco (df_base):", df_base.columns.tolist())
        st.dataframe(df_base.head(6))

        uploaded_file_format = st.file_uploader(
            "Envie o arquivo de notas (Excel/CSV)",
            type=["xlsx", "csv"],
            key="formatar_tab"
        )

        if uploaded_file_format is None:
            st.info("Envie um arquivo para formatar (aba Formatar).")
        else:
            import io
            import re
            import unicodedata

            # ---------- leitura robusta ----------
            try:
                name = str(uploaded_file_format.name).lower()
                if name.endswith('.xlsx'):
                    df_env = pd.read_excel(uploaded_file_format, engine='openpyxl')
                else:
                    try:
                        df_env = pd.read_csv(uploaded_file_format, sep=';')
                    except Exception:
                        try:
                            df_env = pd.read_csv(uploaded_file_format, sep=',')
                        except Exception:
                            uploaded_file_format.seek(0)
                            df_env = pd.read_csv(uploaded_file_format, sep=',', encoding='latin1')
            except Exception as e:
                st.error(f"Erro ao ler o arquivo enviado: {e}")
                df_env = None

            if df_env is None:
                st.warning("Não foi possível ler o arquivo. Verifique o formato/encoding.")
            else:
                st.write("### Diagnóstico rápido (arquivo enviado)")
                st.write("Colunas detectadas:")
                st.write(df_env.columns.tolist())
                st.dataframe(df_env.head(6))

                # ---------- utilidades ----------
                def find_column(df, candidates):
                    cols = df.columns.tolist()
                    for cand in candidates:
                        for c in cols:
                            if c.strip().lower() == cand.strip().lower():
                                return c
                    # busca por substring
                    for cand in candidates:
                        for c in cols:
                            if cand.strip().lower() in c.strip().lower():
                                return c
                    return None

                def normalize_text(s):
                    """Lower, strip, remove accents, collapse spaces."""
                    if pd.isna(s):
                        return ""
                    s = str(s).strip().lower()
                    # normalize unicode (remove accents)
                    s = unicodedata.normalize("NFKD", s)
                    s = "".join(ch for ch in s if not unicodedata.combining(ch))
                    # collapse multiple spaces
                    s = re.sub(r"\s+", " ", s)
                    return s

                def normalize_email(e):
                    if pd.isna(e):
                        return ""
                    e = str(e).strip().lower()
                    # remove stray spaces
                    e = re.sub(r"\s+", "", e)
                    return e

                # ---------- detectar colunas relevantes ----------
                col_nome_env = find_column(df_env, ['Nome completo', 'Nome', 'Aluno', 'Full Name', 'name'])
                col_email_env = find_column(df_env, ['Email institucional', 'Email', 'E-mail', 'email', 'E-mail institucional'])
                col_periodo_env = find_column(df_env, ['Período', 'Periodo', 'Período atual', 'Periodo atual', 'Período letivo'])

                # colunas do banco (conforme você descreveu)
                ra_col_base = find_column(df_base, ['RA', 'Ra', 'ra', 'registro', 'matricula'])
                nome_col_base = find_column(df_base, ['Aluno', 'NOMEALUNO', 'Nome completo', 'Nome', 'ALUNO'])
                email_col_base = find_column(df_base, ['E-mail', 'EMAILALUNO', 'Email institucional', 'Email', 'email'])
                periodo_col_base = find_column(df_base, ['Período letivo', 'Periodo letivo', 'PERIODO', 'Periodo', 'Período'])

                if col_nome_env is None and col_email_env is None:
                    st.error('Não foi possível localizar colunas de nome ou email no arquivo enviado.')
                    st.stop()

                # ---------- cópias de trabalho ----------
                df_env_proc = df_env.copy()
                df_base_proc = df_base.copy()

                # ---------- normalização ----------
                if col_email_env:
                    df_env_proc[col_email_env] = df_env_proc[col_email_env].apply(normalize_email)
                if col_nome_env:
                    df_env_proc[col_nome_env] = df_env_proc[col_nome_env].apply(normalize_text)
                if col_periodo_env:
                    df_env_proc[col_periodo_env] = df_env_proc[col_periodo_env].astype(str).fillna("").str.strip()

                if email_col_base:
                    df_base_proc[email_col_base] = df_base_proc[email_col_base].apply(normalize_email)
                if nome_col_base:
                    df_base_proc[nome_col_base] = df_base_proc[nome_col_base].apply(normalize_text)
                if periodo_col_base:
                    df_base_proc[periodo_col_base] = df_base_proc[periodo_col_base].astype(str).fillna("").str.strip()

                # ---------- MATCH 1: email + periodo (se possível) ----------
                merged_list = []   # guardará os encontrados
                unmatched = df_env_proc.copy()

                if col_email_env and email_col_base and col_periodo_env and periodo_col_base:
                    m1 = unmatched.merge(
                        df_base_proc,
                        left_on=[col_email_env, col_periodo_env],
                        right_on=[email_col_base, periodo_col_base],
                        how='left',
                        suffixes=('_env','_base'),
                        indicator='merge_email_period'
                    )
                    found_m1 = m1[m1['merge_email_period'] == 'both'].copy()
                    if not found_m1.empty:
                        merged_list.append(found_m1)
                    # manter só os que não encontraram
                    unmatched = m1[m1['merge_email_period'] != 'both'].drop(columns=['merge_email_period'])
                # else: pular esta etapa

                # ---------- MATCH 2: email apenas ----------
                if col_email_env and email_col_base and len(unmatched) > 0:
                    m2 = unmatched.merge(
                        df_base_proc,
                        left_on=col_email_env,
                        right_on=email_col_base,
                        how='left',
                        suffixes=('_env','_base'),
                        indicator='merge_email'
                    )
                    found_m2 = m2[m2['merge_email'] == 'both'].copy()
                    if not found_m2.empty:
                        merged_list.append(found_m2)
                    unmatched = m2[m2['merge_email'] != 'both'].drop(columns=[c for c in ['merge_email'] if c in m2.columns])

                # ---------- MATCH 3: nome normalizado (fallback) ----------
                # garantir que col_nome_env e nome_col_base existam
                if col_nome_env and nome_col_base and len(unmatched) > 0:
                    # normalize already applied above
                    m3 = unmatched.merge(
                        df_base_proc,
                        left_on=col_nome_env,
                        right_on=nome_col_base,
                        how='left',
                        suffixes=('_env','_base'),
                        indicator='merge_nome'
                    )
                    found_m3 = m3[m3['merge_nome'] == 'both'].copy()
                    if not found_m3.empty:
                        merged_list.append(found_m3)
                    nao_encontrados = m3[m3['merge_nome'] != 'both'].copy()
                else:
                    # se não puder fazer match por nome, tudo o que sobrou é não encontrado
                    nao_encontrados = unmatched.copy()

                # ---------- concatenar TODOS os encontrados ----------
                if merged_list:
                    df_encontrados = pd.concat(merged_list, ignore_index=True, sort=False)
                else:
                    df_encontrados = pd.DataFrame()

                # garantir df_nao_encontrados definido
                df_nao_encontrados = nao_encontrados if 'nao_encontrados' in locals() else pd.DataFrame()

                # ---------- extrair RA, nome e período (fallback de colunas) ----------
                def get_first_existing(row, candidates):
                    for c in candidates:
                        if c and (c in row) and pd.notna(row[c]) and str(row[c]).strip() != '':
                            return row[c]
                    return ''

                ra_candidates = [ra_col_base, 'RA', 'Ra', 'ra']
                nome_candidates = [nome_col_base, 'Aluno', 'NOMEALUNO', col_nome_env]
                periodo_candidates = [col_periodo_env, periodo_col_base, 'Período letivo', 'Periodo letivo', 'PERIODO']

                # prepara df_encontrados com colunas extraídas
                if not df_encontrados.empty:
                    df_encontrados['RA_extracted'] = df_encontrados.apply(lambda r: get_first_existing(r, ra_candidates), axis=1)
                    df_encontrados['NOME_extracted'] = df_encontrados.apply(lambda r: get_first_existing(r, nome_candidates), axis=1)
                    df_encontrados['PERIODO_extracted'] = df_encontrados.apply(lambda r: get_first_existing(r, periodo_candidates), axis=1)
                else:
                    df_encontrados = pd.DataFrame(columns=['RA_extracted', 'NOME_extracted', 'PERIODO_extracted'])

                # ---------- formatar Student ID (RA) como 7 dígitos sem caracteres ----------
                def fmt_ra(v):
                    try:
                        s = str(v)
                        s = re.sub(r"\D", "", s)
                        return s.zfill(7) if s != '' else ''
                    except Exception:
                        return ''

                df_encontrados['Student ID'] = df_encontrados['RA_extracted'].apply(fmt_ra)

                # dividir nome em First/Last
                # usar nome original do banco em CAIXA ALTA
                df_encontrados['NOME_FINAL'] = df_encontrados[nome_col_base].astype(str).str.upper().str.strip()
                split_nome = df_encontrados['NOME_FINAL'].str.split(n=1, expand=True)
                df_encontrados['First Name'] = split_nome.iloc[:, 0].fillna('') if split_nome.shape[1] >= 1 else ''
                df_encontrados['Last Name'] = split_nome.iloc[:, 1].fillna('') if split_nome.shape[1] > 1 else ''

                # colunas obrigatórias solicitadas
                df_encontrados['Teacher Name'] = ''
                df_encontrados['Gender'] = ''
                df_encontrados['Grade'] = ''

                # Class Name = período (extrair número se existir)
                def normalize_periodo(v):
                    if pd.isna(v) or str(v).strip() == '':
                        return ''
                    s = str(v).strip()
                    m = re.search(r"(\d+)", s)
                    if m:
                        return m.group(1)
                    return s

                df_encontrados['Class Name'] = df_encontrados['PERIODO_extracted'].apply(normalize_periodo)

                # montar df_out na ordem pedida
                cols_required = [
                    'Student ID',
                    'Teacher Name',
                    'First Name',
                    'Last Name',
                    'Gender',
                    'Grade',
                    'Class Name'
                ]
                for c in cols_required:
                    if c not in df_encontrados.columns:
                        df_encontrados[c] = ''

                df_out = df_encontrados[cols_required].copy()

                # ---------- mostrar resultados ----------
                st.subheader('Preview - Alunos encontrados e formatados')
                st.dataframe(df_out.head(200))

                # downloads
                csv_bytes = df_out.to_csv(index=False).encode('utf-8-sig')
                st.download_button('Baixar CSV (formatado)', csv_bytes, file_name='formatado.csv', mime='text/csv')

                xlsx_io = io.BytesIO()
                with pd.ExcelWriter(xlsx_io, engine='openpyxl') as writer:
                    df_out.to_excel(writer, index=False, sheet_name='Formatado')
                xlsx_io.seek(0)
                st.download_button('Baixar XLSX (formatado)', xlsx_io, file_name='formatado.xlsx', mime='application/vnd.openxmlformats-officedocument-spreadsheetml.sheet')

                # ---------- Mostrar não encontrados ----------
                st.subheader('Alunos NÃO encontrados no banco (precisa revisar)')
                if not df_nao_encontrados.empty:
                    cols_show = []
                    if col_nome_env:
                        cols_show.append(col_nome_env)
                    if col_email_env:
                        cols_show.append(col_email_env)
                    if col_periodo_env:
                        cols_show.append(col_periodo_env)
                    # evitar erro se cols_show vazia
                    cols_show = [c for c in cols_show if c in df_nao_encontrados.columns]
                    if cols_show:
                        st.dataframe(df_nao_encontrados[cols_show].drop_duplicates().head(200))
                        csv_nf = df_nao_encontrados[cols_show].drop_duplicates().to_csv(index=False).encode('utf-8-sig')
                        st.download_button('Baixar NÃO encontrados (CSV)', csv_nf, file_name='nao_encontrados.csv', mime='text/csv')
                    else:
                        st.write(df_nao_encontrados.head(200))
                else:
                    st.write('Todos os registros foram encontrados por email ou nome.')

                # resumo simples
                st.write(f"Linhas no arquivo: {len(df_env)}")
                st.write(f"Encontrados: {len(df_out)}")
                st.write(f"Não encontrados: {len(df_nao_encontrados)}")