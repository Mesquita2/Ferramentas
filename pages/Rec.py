from datetime import date
import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.shared import Inches
import streamlit as st
import pandas as pd
import io
from auth import check_authentication

imagem_rodape = "Endereço.jpeg"
imagem_cabecalho = 'Logo.jpg'
ARQUIVOBASE = "alunosxdisciplinas"
ARQUIVOREC = "rec"


# Configuração da página
st.set_page_config(page_title="Limpeza Dados da REC", 
                   page_icon=" ", 
                   layout="wide")

if not check_authentication():
    st.stop()

# Função para substituir o arquivo de alunos
def limpar_rec(df):
    if df is not None:
        df_base = st.session_state["dados"].get(ARQUIVOBASE).copy()
        df['DISCIPLINA'] = (
            df['DISCIPLINA']
            .str.replace(r'\s*\([^()]*\)\s*$', '', regex=True)  # remove apenas o último parêntese
            .str.replace(r'[\u200b\u200e\u202c\u00a0]', '', regex=True) 
            .str.strip()
        )
        
        
        df["RA"] = df["RA"].astype(str).str.zfill(7)
        df_base["RA"] = df_base["RA"].astype(str).str.zfill(7)
        
        df.rename(columns={'VALOR': 'DISCIPLINA',
                            'RA': 'RA'}, inplace=True)
        
        df = pd.merge(df, df_base[['DISCIPLINA', 'RA',  'TURMADISC', 'ALUNO']],
                  on=['DISCIPLINA', 'RA'],
                  how='left')        

        df = df.drop_duplicates(subset=['ALUNO', 'DISCIPLINA', 'TURMADISC', 'RA'])
        
        df = df[df['CODSTATUS'] != 'C']
        
        df['RA'] = df['RA'].apply(lambda x: str(x).zfill(7))
        st.success("Dados de alunos substituídos com sucesso!")
        return df
    else:
        st.warning("Não existe arquivo REC, Voltar a pagina Inicial!")
        
def adicionar_imagem_no_cabecalho(doc, imagem_cabecalho):
    # Acessando o cabeçalho da primeira seção do documento
    section = doc.sections[0]
    header = section.header

    # Criando um parágrafo no cabeçalho e adicionando uma imagem
    paragraph = header.paragraphs[0]  # Usando o primeiro parágrafo do cabeçalho
    section.header_distance = Inches(0.2)   
    run = paragraph.add_run()

    # Adicionando a imagem ao cabeçalho
    run.add_picture(imagem_cabecalho, width=Inches(7.5), height=Inches(1))  # Ajuste o tamanho conforme necessário

def adicionar_imagem_no_rodape(doc, imagem_rodape):
    # Acessando o rodapé da primeira seção do documento
    section = doc.sections[0]
    footer = section.footer
    

    # Criando um parágrafo no rodapé e adicionando uma imagem
    paragraph = footer.paragraphs[0]  # Usando o primeiro parágrafo do rodapé
    section.footer_distance = Inches(0.2)
    run = paragraph.add_run()

    # Adicionando a imagem ao rodapé
    run.add_picture(imagem_rodape, width=Inches(7.5), height=Inches(1))  # Ajuste o tamanho conforme necessário


def gerar_relatorio(df, disciplinas, turmas):
    dataatual = date.today().strftime('%d/%m/%Y')
    df = df[df["DISCIPLINA"].isin(disciplinas) & df["TURMADISC"].isin(turmas)]
    df = df.sort_values(by=["DISCIPLINA", "TURMADISC", "ALUNO"])

    doc = Document()
    adicionar_imagem_no_cabecalho(doc, imagem_cabecalho)
    adicionar_imagem_no_rodape(doc, imagem_rodape)

    section = doc.sections[0]
    section.left_margin = Inches(0.5) 
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.5)  
    section.bottom_margin = Inches(0.5) 

    for (disciplina, turma), df_grupo in df.groupby(['DISCIPLINA', 'TURMADISC']):
        # Título
        p = doc.add_paragraph()
        p.add_run("\n\n")
        run = p.add_run(f"Disciplina: {disciplina}")
        run.font.name = 'Arial'
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0, 0, 0)

        p = doc.add_paragraph()
        run = p.add_run(f"Turma: {turma}")
        run.font.name = 'Arial'
        run.font.size = Pt(12)

        p = doc.add_paragraph()
        run = p.add_run(f"Data: {dataatual}")
        run.font.name = 'Arial'
        run.font.size = Pt(12)

        # Tabela
        df_grupo = df_grupo[['ALUNO']].copy()
        df_grupo['ASSINATURA'] = ''
        table = doc.add_table(rows=1, cols=len(df_grupo.columns))
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        for i, col_name in enumerate(df_grupo.columns):
            hdr_cells[i].text = col_name

        for _, row in df_grupo.iterrows():
            row_cells = table.add_row().cells
            for i, item in enumerate(row):
                row_cells[i].text = str(item)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def dash(df):
    if not df:
        st.write("Data frame Vazio")
        return pd.DataFrame() 
    if not os.path.exists(df):
        st.write(f"Erro: Arquivo '{df}' não encontrado.")
        return pd.DataFrame()  
    return pd.read_excel(df) 

# Interface do Streamlit
st.title("Limpeza e tratamento de notas de REC")
        
st.subheader("Dados dos Cadastrados na REC")
df_cadastro = st.session_state["dados"].get(ARQUIVOREC).copy()
df = df_cadastro.copy()
if df_cadastro is not None: 
    st.dataframe(df_cadastro[['DISCIPLINA', 'NOME']])


def gerar_excel(df_rec, disciplinas, turmas):
    df_filtrado = df_rec[df_rec["DISCIPLINA"].isin(disciplinas) & df_rec["TURMADISC"].isin(turmas)].copy()
    df_filtrado['RA'] = df_filtrado['RA'].astype(str).str.zfill(7)
    df_filtrado['NOTAS'] = 0
    colunas = ['TURMADISC', 'DISCIPLINA', 'RA', 'ALUNO', 'NOTAS']
    df_filtrado = df_filtrado[colunas].sort_values(by=["DISCIPLINA", "TURMADISC", "ALUNO"])
    
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
        df_filtrado.to_excel(writer, index=False, sheet_name="Notas")
    output.seek(0)
    return output


st.title("Gerador de Planilha de Notas para REC")

df_rec = limpar_rec(df)
if df_rec.empty:
    st.stop()
    
df_filtrado = df[df['NOME'] == 'GUILHERME MÁXIMUS MOTA LOPES']
st.write(df_filtrado.columns)
st.write("**Dados da REC para RA 1414293:**")
st.dataframe(df_filtrado[['DISCIPLINA', 'TURMADISC', 'NOME', 'RA']])    
        
    
disciplinas = df_rec["DISCIPLINA"].unique().tolist()
disciplinas_selecionadas = st.multiselect("1. Escolha as disciplinas", disciplinas)

if disciplinas_selecionadas:
    turmas_disponiveis = df_rec[df_rec["DISCIPLINA"].isin(disciplinas_selecionadas)]["TURMADISC"].unique().tolist()
    turmas_selecionadas = st.multiselect("2. Escolha as turmas", turmas_disponiveis)

    if turmas_selecionadas:
        prova = st.selectbox("3. Escolha se é REC_P1 ou REC_P2 ou REC_FINAL", ["REC_P1", "REC_P2", "REC_FINAL"])

        # Filtra os dados para visualização
        df_filtrado = df_rec[
            (df_rec["DISCIPLINA"].isin(disciplinas_selecionadas)) &
            (df_rec["TURMADISC"].isin(turmas_selecionadas))
        ]

        st.write(f"**Alunos da(s) Disciplina(s): {disciplinas_selecionadas} | Turma(s): {turmas_selecionadas}**")
        total = df_filtrado['ALUNO'].count()
        st.write(f"**Quantidade de REC solicitadas: {total}**")
        df_filtrado = df_filtrado.sort_values(by="ALUNO", ascending=True)
        st.dataframe(df_filtrado[["ALUNO", "DISCIPLINA", "TURMADISC"]])

        # Botão para gerar planilha Excel
        excel_file = gerar_excel(df_rec, disciplinas_selecionadas, turmas_selecionadas)
        st.download_button(
            label="⬇ Gerar e Baixar Planilha Excel (Multi)",
            data=excel_file,
            file_name=f"Planilha_REC_{prova}.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )

        st.title("Criar Relatório de Assinatura")
        relatorio = gerar_relatorio(df_rec, disciplinas_selecionadas, turmas_selecionadas)
        st.download_button(
            label="📄 Gerar e Baixar Relatório de Assinaturas",
            data=relatorio,
            file_name=f"Relatorio_Assinatura_REC_{prova}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

