from datetime import date
import pandas as pd
import io
from docx import Document
from docx.shared import Pt, RGBColor, Inches
import streamlit as st


# Função de transformação dos eventos
def transformar_eventos(df):
    colunas_eventos = [col for col in df.columns if 'Eventos' in col]
    for col in colunas_eventos:
        df[col] = df[col].fillna('').apply(lambda x: [item.strip() for item in x.split(';') if item.strip() != ''])
    for col in colunas_eventos:
        df = df.explode(col)
    df_long = df.melt(
        id_vars=['Carimbo de data/hora', 'Nome Completo', 'e-mail', 'Whatsapp'],
        value_vars=colunas_eventos,
        var_name='Horário',
        value_name='Palestra'
    )
    df_long = df_long[df_long['Palestra'] != '']
    df_long = df_long.dropna(subset=['Palestra'])
    return df_long

# Função de gerar o relatório por palestra
def gerar_relatorio_palestra(df, palestra, imagem_cabecalho, imagem_rodape):
    doc = Document()
    
    # Ajusta as margens
    section = doc.sections[0]
    section.left_margin = Inches(0.5) 
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.5)  
    section.bottom_margin = Inches(0.5)
    
    # Cabeçalho com imagem
    section = doc.sections[0]
    header = section.header
    paragraph = header.paragraphs[0]
    section.header_distance = Inches(0.2)
    run = paragraph.add_run()
    run.add_picture(imagem_cabecalho, width=Inches(7.5), height=Inches(1))
    
    # Rodapé com imagem
    footer = section.footer
    paragraph = footer.paragraphs[0]
    section.footer_distance = Inches(0.2)
    run = paragraph.add_run()
    run.add_picture(imagem_rodape, width=Inches(7.5), height=Inches(1))
    dataatual = date.today().strftime('%d/%m/%Y')
    
    # Título
    p = doc.add_paragraph()
    run = p.add_run(f"Lista de Presença - {palestra}\nData:{dataatual}")
    run.font.name = 'Arial'
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0, 0, 0)
    
    # Tabela sem e-mail
    df = df[['Nome Completo']].copy()
    df['ASSINATURA'] = '  '
    
    df['Nome Completo'] = df['Nome Completo'].str.title()
    df = df.sort_values('Nome Completo', key=lambda col: col.str.normalize('NFKD'))
    
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, col_name in enumerate(df.columns):
        hdr_cells[i].text = col_name
    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, item in enumerate(row):
            row_cells[i].text = str(item)
    
    # Salva em memória
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# Streamlit interface
# Streamlit interface
st.title("Gerar Relatório de Assinaturas por Palestra")

uploaded_file = st.file_uploader("Escolha o arquivo Excel com os eventos", type=['xlsx'])

imagem_cabecalho = 'Logo.jpg'
imagem_rodape = 'Endereço.jpeg'

if uploaded_file is not None:
    df_eventos = pd.read_excel(uploaded_file)
    df_organizado = transformar_eventos(df_eventos)
    
    st.write("Dados Organizados:")
    st.dataframe(df_organizado)
    
    palestras_unicas = df_organizado['Palestra'].unique().tolist()
    
    palestra_selecionada = st.selectbox("Selecione a palestra para gerar o relatório:", palestras_unicas)
    
    if palestra_selecionada:
        df_palestra = df_organizado[df_organizado['Palestra'] == palestra_selecionada]
        relatorio = gerar_relatorio_palestra(df_palestra, palestra_selecionada, imagem_cabecalho, imagem_rodape)
        
        st.download_button(
            label=f"⬇ Baixar Relatório: {palestra_selecionada}",
            data=relatorio,
            file_name=f"Relatorio_{palestra_selecionada}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
