from datetime import date
import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.shared import Inches
import streamlit as st
import pandas as pd
import io
from auth import check_authentication

imagem_rodape = "Endereço.jpeg"
imagem_cabecalho = 'Logo.jpg'
ARQUIVOBASE = "alunosxdisciplinas"
ARQUIVOREC = "rec"

# Configuração da página
st.set_page_config(page_title="Limpeza Dados da REC", 
                   page_icon=" ", 
                   layout="wide")

if not check_authentication():
    st.stop()

# Função para limpar e preparar dados
def limpar_rec(df):
    if df is not None:
        df["RA"] = df["RA"].astype(str).str.zfill(7)
        df.rename(columns={'VALOR': 'DISCIPLINA', 'RA': 'RA'}, inplace=True)
        df = df[df['NOMESTATUS'] == 'Período em Curso']
        df['RA'] = df['RA'].apply(lambda x: str(x).zfill(7))
        st.success("Dados de alunos substituídos com sucesso!")
        return df
    else:
        st.warning("Não existe arquivo REC, Voltar à página Inicial!")

# Excel geral de alunos em período em curso
def gerar_excel_em_curso(df):
    df_export = df[['RA', 'ALUNO', 'DISCIPLINA', 'TURMADISC', 'NOMESTATUS']].copy()
    df_export = df_export.sort_values(by=["DISCIPLINA", "TURMADISC", "ALUNO"])
    
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
        df_export.to_excel(writer, index=False, sheet_name="Em_Periodo")
    output.seek(0)
    return output

# Excel com filtros
def gerar_excel_com_filtros(df_rec, disciplinas, turmas):
    df_filtrado = df_rec[df_rec["DISCIPLINA"].isin(disciplinas) & df_rec["TURMADISC"].isin(turmas)].copy()
    df_filtrado['RA'] = df_filtrado['RA'].astype(str).str.zfill(7)
    df_filtrado['NOTAS'] = 0
    colunas = ['TURMADISC', 'DISCIPLINA', 'RA', 'ALUNO', 'NOTAS']
    df_filtrado = df_filtrado[colunas].sort_values(by=["DISCIPLINA", "TURMADISC", "ALUNO"])
    
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
        df_filtrado.to_excel(writer, index=False, sheet_name="Notas")
    output.seek(0)
    return output

def gerar_relatorio_assinatura(df, disciplinas, turmas):
    data_hoje = date.today().strftime("%d/%m/%Y")
    df = df[(df["DISCIPLINA"].isin(disciplinas)) & (df["TURMADISC"].isin(turmas))].copy()
    df = df.sort_values(by=["DISCIPLINA", "TURMADISC", "ALUNO"])

    doc = Document()

    # Definições de margem
    section = doc.sections[0]
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)

    for (disciplina, turma), grupo in df.groupby(["DISCIPLINA", "TURMADISC"]):
        doc.add_paragraph(f"Disciplina: {disciplina}", style='Heading 2')
        doc.add_paragraph(f"Turma: {turma}")
        doc.add_paragraph(f"Data: {data_hoje}")
        doc.add_paragraph(" ")

        tabela = doc.add_table(rows=1, cols=2)
        tabela.style = "Table Grid"
        tabela.autofit = True

        hdr_cells = tabela.rows[0].cells
        hdr_cells[0].text = 'Aluno'
        hdr_cells[1].text = 'Assinatura'

        for _, row in grupo.iterrows():
            linha = tabela.add_row().cells
            linha[0].text = row["ALUNO"]
            linha[1].text = " "

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output


# Interface Streamlit
st.title("Limpeza e tratamento de notas de REC")

df_cadastro = st.session_state["dados"].get(ARQUIVOBASE).copy()
df = df_cadastro.copy()

df.rename(columns={'NOMEDISCIPLINA': 'DISCIPLINA',
                   'NOMECURSO': 'CURSO',
                   'NOMEALUNO': 'ALUNO'}, inplace=True)

st.subheader("Dados dos Cadastrados na REC")

df_rec = limpar_rec(df)
if df_rec.empty:
    st.stop()

# 🔍 Relatório geral
st.subheader(" Relatório: Alunos em Período em Curso")
df_em_curso = df_rec.copy()
if not df_em_curso.empty:
    st.dataframe(df_em_curso[['RA', 'ALUNO', 'DISCIPLINA', 'TURMADISC','NOMESTATUS']])
    total = df_em_curso['RA'].nunique()
    st.info(f"Total de alunos em 'Período em Curso': **{total}**")
    
# Gerador com filtros aprimorados
st.title("Gerador de Planilha de Notas para REC")

# Filtro por texto para disciplinas
disciplinas = df_rec["DISCIPLINA"].dropna().unique().tolist()
disciplinas_selecionadas = st.multiselect("1. Escolha as disciplinas", disciplinas)

if disciplinas_selecionadas:
    # Filtro por texto para turmas
    turmas_disponiveis = df_rec[df_rec["DISCIPLINA"].isin(disciplinas_selecionadas)]["TURMADISC"].dropna().unique().tolist()
    turmas_selecionadas = st.multiselect("2. Escolha as turmas", turmas_disponiveis)

    if turmas_selecionadas:

        df_filtrado = df_rec[
            (df_rec["DISCIPLINA"].isin(disciplinas_selecionadas)) &
            (df_rec["TURMADISC"].isin(turmas_selecionadas))
        ]

        st.write(f"**Alunos da(s) Disciplina(s): {disciplinas_selecionadas} | Turma(s): {turmas_selecionadas}**")
        st.write(f"**Quantidade de REC solicitadas: {df_filtrado['ALUNO'].count()}**")
        st.write(f"**Quantidade de alunos distintos: {df_filtrado['ALUNO'].nunique()}**")
        st.dataframe(df_filtrado[["ALUNO", "DISCIPLINA", "TURMADISC"]])

        curso = df_filtrado['CURSO'].iloc[0]
        
                # Botão para gerar o relatório de assinaturas
        relatorio_docx = gerar_relatorio_assinatura(df_rec, disciplinas_selecionadas, turmas_selecionadas)
        st.download_button(
            label="Gerar Relatório para Impressão",
            data=relatorio_docx,
            file_name=f"Relatorio_Assinaturas_{curso}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        
        relatorio_excel = gerar_excel_com_filtros(df_rec, disciplinas_selecionadas, turmas_selecionadas)
        st.download_button(
            label="Gerar Relatório de Notas",
            data=relatorio_excel,
            file_name=f"Relatorio_Notas_{curso}.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )
        
        

