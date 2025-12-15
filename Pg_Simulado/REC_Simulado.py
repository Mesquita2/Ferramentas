from datetime import date
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
import streamlit as st
import pandas as pd
import io

imagem_rodape = "Endereço.jpeg"
imagem_cabecalho = 'Logo.jpg'

ARQUIVOBASE = "alunosxdisciplinas"
ARQUIVOREC = "rec_simulado"


def carregar():

    # --------------------------
    # LIMPA E COMPLETA O REC
    # --------------------------
    def limpar_rec(df):
        if df is None:
            st.warning("Não existe arquivo REC, volte à página Inicial!")
            return pd.DataFrame()

        df_base = st.session_state["dados"].get(ARQUIVOBASE).copy()

        # Padroniza RA
        df["RA"] = df["RA"].astype(str).str.zfill(7)
        df_base["RA"] = df_base["RA"].astype(str).str.zfill(7)

        # Renomear NOME -> ALUNO (se existir)
        if "NOME" in df.columns:
            df.rename(columns={'NOME': 'ALUNO'}, inplace=True)

        # Remove duplicados da planilha REC
        df = df.drop_duplicates(subset=['RA'])

        # MERGE para buscar turma
        df = df.merge(
            df_base[['RA', 'CODTURMA']],
            on='RA',
            how='left'
        )

        # 🚫 Remove duplicatas causadas pelo df_base (caso um RA apareça mais de 1 vez na base)
        df = df.drop_duplicates(subset=['RA'], keep='first')

        # Identifica RA sem turma
        sem_turma = df[df["CODTURMA"].isna()]
        if not sem_turma.empty:
            st.error("⚠ Alguns alunos não tiveram turma encontrada na base!")
            st.dataframe(sem_turma)
            st.stop()

        return df

    # --------------------------
    # CABEÇALHO E RODAPÉ
    # --------------------------
    def adicionar_imagem_no_cabecalho(doc, imagem_cabecalho):
        section = doc.sections[0]
        header = section.header
        paragraph = header.paragraphs[0]
        section.header_distance = Inches(0.2)
        run = paragraph.add_run()
        run.add_picture(imagem_cabecalho, width=Inches(7.5), height=Inches(1))

    def adicionar_imagem_no_rodape(doc, imagem_rodape):
        section = doc.sections[0]
        footer = section.footer
        paragraph = footer.paragraphs[0]
        section.footer_distance = Inches(0.2)
        run = paragraph.add_run()
        run.add_picture(imagem_rodape, width=Inches(7.5), height=Inches(1))

    # --------------------------
    # GERAR RELATÓRIO DE ASSINATURA
    # --------------------------
    def gerar_relatorio(df_filtrado, turmas_escolhidas):
        df = df_filtrado.copy()
        dataatual = date.today().strftime('%d/%m/%Y')

        doc = Document()
        adicionar_imagem_no_cabecalho(doc, imagem_cabecalho)
        adicionar_imagem_no_rodape(doc, imagem_rodape)

        sec = doc.sections[0]
        for side in ("left", "right", "top", "bottom"):
            setattr(sec, f"{side}_margin", Inches(0.5))

        p = doc.add_paragraph()
        run = p.add_run("\n\nSIMULADO DE RECUPERAÇÃO")
        run.font.name = 'Arial'
        run.font.size = Pt(14)

        p = doc.add_paragraph()
        run = p.add_run(f"Turmas: {', '.join(turmas_escolhidas)}")
        run.font.name = 'Arial'
        run.font.size = Pt(12)

        p = doc.add_paragraph()
        run = p.add_run(f"Data: {dataatual}")
        run.font.name = 'Arial'
        run.font.size = Pt(12)

        tabela = doc.add_table(rows=1, cols=2)
        tabela.style = 'Table Grid'
        hdr = tabela.rows[0].cells
        hdr[0].text = 'ALUNO'
        hdr[1].text = 'ASSINATURA'

        for _, row in df[['ALUNO']].iterrows():
            cells = tabela.add_row().cells
            cells[0].text = str(row['ALUNO'])
            cells[1].text = ''

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf

    # --------------------------
    # GERAR PLANILHA EXCEL
    # --------------------------
    def gerar_excel(df_filtrado):
        df = df_filtrado.copy()
        df["RA"] = df["RA"].astype(str)
        df["NOTAS"] = 0
        df = df[['CODTURMA', 'RA', 'ALUNO', 'NOTAS']].sort_values('ALUNO')

        out = io.BytesIO()
        with pd.ExcelWriter(out, engine='xlsxwriter') as writer:
            df.to_excel(writer, index=False, sheet_name="Notas")
        out.seek(0)
        return out

    # --------------------------
    # INTERFACE STREAMLIT
    # --------------------------
    st.title("Limpeza e Tratamento de NOTAS da REC")

    df_cadastro = st.session_state["dados"].get(ARQUIVOREC)

    if df_cadastro is not None:
        st.subheader("Arquivo REC recebido (RA + Nome)")
        colunas_validas = [c for c in ["RA", "NOME", "ALUNO"] if c in df_cadastro.columns]
        st.dataframe(df_cadastro[colunas_validas])
    else:
        st.warning("Nenhum arquivo REC carregado!")
        st.stop()

    # Faz limpeza e inclui turma automaticamente
    df_rec = limpar_rec(df_cadastro.copy())

    turmas = sorted(df_rec["CODTURMA"].unique().tolist())

    # MULTISELECT DE TURMAS
    turmas_selecionadas = st.multiselect("Escolha as Turmas", turmas)

    if not turmas_selecionadas:
        st.warning("Selecione pelo menos uma turma!")
        st.stop()

    prova = st.selectbox("Escolha a Prova", ["REC_P3"])

    # Filtra várias turmas
    df_filtrado = df_rec[df_rec["CODTURMA"].isin(turmas_selecionadas)].sort_values("ALUNO")

    total = len(df_filtrado)

    st.markdown(f"**Turmas selecionadas: {', '.join(turmas_selecionadas)}**")
    st.markdown(f"**Total de REC solicitadas: {total}**")
    st.dataframe(df_filtrado[["RA", "ALUNO", "CODTURMA"]])

    # --------------------------
    # BOTÕES DE DOWNLOAD
    # --------------------------
    excel = gerar_excel(df_filtrado)
    st.download_button(
        "⬇ Baixar Planilha Excel",
        data=excel,
        file_name=f"{'-'.join(turmas_selecionadas)}_{prova}.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )

    relatorio = gerar_relatorio(df_filtrado, turmas_selecionadas)
    st.download_button(
        "⬇ Baixar Relatório de Assinaturas",
        data=relatorio,
        file_name=f"{'-'.join(turmas_selecionadas)}_{prova}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )